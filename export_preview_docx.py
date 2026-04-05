# -*- coding: utf-8 -*-
"""匯出預覽為 .docx（需安裝 python-docx）。

- save_screen_layout_docx：排列為「與篩選區塊相同」時，以表格重現多欄流式排版，拋棄式姓名加藍色字元框線（同主篩選外框）。
- save_preview_text_as_docx：由純文字分段落＋樣式規則轉換。
- save_paragraph_runs_docx：由（段落 → 帶 bold/italic/underline 的 run）寫入，供編輯區套用。
"""

from __future__ import annotations

import re
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_UNDERLINE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

FONT_EA = "Microsoft JhengHei"
FONT_LATIN = "Calibri"
TITLE_RGB = RGBColor(0x0D, 0x47, 0xA1)
FOOTER_RGB = RGBColor(0x1A, 0x23, 0x7E)
# 主篩選拋棄式外框色（與 app FILTER 畫布 outline 一致）
DISPOSABLE_BORDER_COLOR = "0D47A1"

_DISP = re.compile(r"(〖[^〗]*〗)")


def _set_run_east_asia_base(run, *, size_pt: float = 12) -> None:
    run.font.name = FONT_LATIN
    run.font.size = Pt(size_pt)
    rfonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    rfonts.set(qn("w:eastAsia"), FONT_EA)


def _set_run_font(
    run,
    *,
    size_pt: float = 12,
    bold: bool = False,
    color: RGBColor | None = None,
) -> None:
    _set_run_east_asia_base(run, size_pt=size_pt)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def save_paragraph_runs_docx(
    path: str | Path,
    paragraphs: list[list[tuple[str, dict[str, bool]]]],
) -> None:
    """由（段落 → 連續 run 與樣式）寫入 .docx；樣式鍵：bold, italic, underline。"""
    path = Path(path)
    doc = Document()
    for runs in paragraphs:
        p = doc.add_paragraph()
        if not runs:
            continue
        for text, st in runs:
            if not text:
                continue
            run = p.add_run(text)
            _set_run_east_asia_base(run)
            run.bold = bool(st.get("bold"))
            run.italic = bool(st.get("italic"))
            run.underline = (
                WD_UNDERLINE.SINGLE if st.get("underline") else WD_UNDERLINE.NONE
            )
    doc.save(str(path))


def _add_run_char_border(run, *, color_hex: str = DISPOSABLE_BORDER_COLOR, sz_eighths_pt: str = "8") -> None:
    """Word 字元框線（對應主篩選 Canvas 藍框，中空不加底色）。"""
    rpr = run._element.get_or_add_rPr()
    bdr = OxmlElement("w:bdr")
    bdr.set(qn("w:val"), "single")
    bdr.set(qn("w:sz"), sz_eighths_pt)
    bdr.set(qn("w:space"), "1")
    bdr.set(qn("w:color"), color_hex)
    bdr.set(qn("w:frame"), "1")
    rpr.append(bdr)


def _set_table_borders_nil(table) -> None:
    tbl = table._element
    tbl_pr = tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        tbl.insert(0, tbl_pr)
    tbl_borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "nil")
        el.set(qn("w:sz"), "0")
        el.set(qn("w:space"), "0")
        tbl_borders.append(el)
    tbl_pr.append(tbl_borders)


def _fill_cell_with_roster(paragraph, app: Any, r: Any, rule: dict) -> None:
    paragraph.text = ""
    paragraph.paragraph_format.space_after = Pt(0)
    segs = app._roster_segments(r, rule)
    first = True
    for text, framed in segs:
        if not text:
            continue
        if not first:
            paragraph.add_run(" ")
        first = False
        run = paragraph.add_run(text)
        if framed:
            _set_run_font(run)
            _add_run_char_border(run)
        else:
            _set_run_font(run)


def _is_block_title_line(line: str) -> bool:
    t = line.strip()
    if not t or "\t" in t:
        return False
    return "（" in t and t.endswith("筆）")


def _is_stat_or_footer_line(line: str) -> bool:
    t = line.strip()
    return (
        t.startswith("統計")
        or t.startswith("【全部")
        or t.startswith("【本次")
    )


def save_preview_text_as_docx(path: str | Path, text: str) -> None:
    path = Path(path)
    doc = Document()

    lines = text.replace("\r\n", "\n").replace("\r", "\n").split("\n")
    for line in lines:
        if line == "":
            doc.add_paragraph()
            continue

        title = _is_block_title_line(line)
        stat = _is_stat_or_footer_line(line) and not title
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)

        if title:
            run = p.add_run(line)
            _set_run_font(run, size_pt=13, bold=True, color=TITLE_RGB)
            continue

        if stat:
            run = p.add_run(line)
            col = FOOTER_RGB if line.strip().startswith("【") else TITLE_RGB
            _set_run_font(run, bold=True, color=col)
            continue

        parts = _DISP.split(line)
        for part in parts:
            if not part:
                continue
            run = p.add_run(part)
            is_disp = part.startswith("〖") and part.endswith("〗")
            if is_disp:
                inner = part[1:-1]
                run.text = inner
                _set_run_font(run)
                _add_run_char_border(run)
            else:
                _set_run_font(run)

    doc.save(str(path))


def save_screen_layout_docx(
    path: str | Path,
    app: Any,
    tags_subset: list[str],
    *,
    include_block_stats: bool,
    include_footer: bool,
) -> None:
    """依目前資料與「與篩選區塊相同」邏輯產生多欄表格；不依賴預覽文字框內容。"""
    from tag_store import list_hashtags

    path = Path(path)
    doc = Document()
    lib_list = list_hashtags()
    order_map = {name: i for i, name in enumerate(lib_list)}
    tags = sorted(tags_subset, key=lambda t: order_map.get(t, 10**9))

    pitch_basis = app._visible_primary_filter_tags() or list(tags_subset)
    unified_pitch = app._filter_unified_roster_pitch_for_plain_export(pitch_basis)
    content_w = app._export_filter_block_content_width_px()

    sec = doc.sections[0]
    try:
        uw_in = float(sec.page_width.inches) - float(sec.left_margin.inches) - float(sec.right_margin.inches)
    except (AttributeError, TypeError, ValueError):
        uw_in = 6.5

    title_needs_extra_top = False
    for tag in tags:
        matches = app._rows_matching_tag_value(tag)
        if not matches:
            continue
        rule = app._get_display_rule(tag)

        p0 = doc.add_paragraph()
        if title_needs_extra_top:
            p0.paragraph_format.space_before = Pt(10)
            title_needs_extra_top = False
        r0 = p0.add_run(tag)
        _set_run_font(r0, size_pt=13, bold=True, color=TITLE_RGB)
        p0.paragraph_format.space_after = Pt(4)

        groups = app._export_filter_roster_table_row_groups(
            matches, rule, unified_pitch, content_w
        )
        groups = [g for g in groups if g]
        if not groups:
            title_needs_extra_top = True
            continue

        ncols = max(len(g) for g in groups)
        table = doc.add_table(rows=len(groups), cols=ncols)
        _set_table_borders_nil(table)
        col_w = Inches(max(uw_in / max(ncols, 1), 0.85))

        for ci in range(ncols):
            try:
                table.columns[ci].width = col_w
            except (IndexError, AttributeError):
                break

        for ri, grp in enumerate(groups):
            row_cells = table.rows[ri].cells
            for ci in range(ncols):
                cell = row_cells[ci]
                cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
                para = cell.paragraphs[0]
                if ci < len(grp):
                    _fill_cell_with_roster(para, app, grp[ci], rule)
                else:
                    para.text = ""

        if include_block_stats:
            stat_lines = app._primary_filter_block_stats_lines(matches, rule)
            for si, st in enumerate(stat_lines):
                ps = doc.add_paragraph()
                ps.paragraph_format.space_before = Pt(0)
                ps.paragraph_format.space_after = Pt(14 if si == len(stat_lines) - 1 else 2)
                rs = ps.add_run(st)
                _set_run_font(rs, bold=True, color=TITLE_RGB)
        else:
            title_needs_extra_top = True

    if include_footer and tags:
        vis_foot = app._visible_primary_filter_tags()
        foot = app._export_footer_text_for_tags(
            tags,
            visible_tags=vis_foot if vis_foot else None,
        )
        for fi, line in enumerate(foot.split("\n")):
            pf = doc.add_paragraph()
            rf = pf.add_run(line)
            _set_run_font(rf, bold=True, color=FOOTER_RGB)
            pf.paragraph_format.space_before = Pt(6 if fi == 0 else 2)

    doc.save(str(path))
